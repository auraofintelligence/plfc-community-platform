from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = ROOT / "templates"
TEMPLATES.mkdir(exist_ok=True)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
styles["Normal"].font.name = "Atkinson Hyperlegible"
styles["Normal"].font.size = Pt(10.5)
styles["Heading 1"].font.name = "Nunito Sans"
styles["Heading 1"].font.size = Pt(22)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = RGBColor(7, 68, 95)
styles["Heading 2"].font.name = "Nunito Sans"
styles["Heading 2"].font.size = Pt(15)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = RGBColor(4, 127, 168)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Point Lookout Fishing Club\nBackend Governance Starter Pack")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(7, 68, 95)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Drive, Sheets, Gmail, Calendar, agents and public-safe exports").italic = True

doc.add_paragraph(
    "Purpose: give the committee a plain operating model for using Google Drive as the private vault while the public website only receives approved summaries."
)

doc.add_heading("1. Core Rule", level=1)
core = doc.add_paragraph()
core.add_run("Raw data goes in privately. Approved meaning comes out publicly.").bold = True
doc.add_paragraph(
    "Members, payments, precise locations, emergency contacts, private media and meeting drafts stay in controlled club systems. Public pages show broad zones, counts, status and reviewed stories."
)

doc.add_heading("2. Storage Model", level=1)
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = ["Layer", "Tool", "Contains", "Public rule"]
for cell, text in zip(table.rows[0].cells, headers):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for row in [
    ["Vault", "Google Drive", "Original files, uploads, minutes, media", "Private by default"],
    ["Ledgers", "Google Sheets", "Members, payments, permissions, events", "Export summaries only"],
    ["Messages", "Gmail", "Inbound requests and outbound confirmations", "Do not publish raw threads"],
    ["Schedule", "Google Calendar", "Committee, comps, working bees", "Create only after approval"],
    ["Public site", "GitHub Pages", "Approved stories and dashboards", "No private data"],
]:
    cells = table.add_row().cells
    for cell, text in zip(cells, row):
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.add_heading("3. Suggested Drive Folders", level=1)
for item in [
    "PLFC Operations Vault / 00 Inbox to sort",
    "PLFC Operations Vault / 01 Members private",
    "PLFC Operations Vault / 02 Competitions and field ops",
    "PLFC Operations Vault / 03 Meetings and minutes",
    "PLFC Operations Vault / 04 Finance and Stripe",
    "PLFC Operations Vault / 05 Media permissions",
    "PLFC Operations Vault / 06 Grants and sponsors",
    "PLFC Operations Vault / 99 Public-safe exports",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4. Promptness Rules", level=1)
prompt_table = doc.add_table(rows=1, cols=4)
prompt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
prompt_table.style = "Table Grid"
for cell, text in zip(prompt_table.rows[0].cells, ["Timing", "Data", "Receiver", "Rule"]):
    cell.text = text
for row in [
    ["Immediate", "Event check-ins, weather flags", "Field lead", "Operational only"],
    ["Daily", "New members, payments, media queue", "Secretary / Treasurer", "Digest"],
    ["Weekly", "Actions, reports, meeting pack", "Committee", "Review pack"],
    ["After review", "Photos, routes, public stories", "Public", "Permission checked"],
]:
    cells = prompt_table.add_row().cells
    for cell, text in zip(cells, row):
        cell.text = text

doc.add_heading("5. Journey Context", level=1)
doc.add_paragraph(
    "The wrap-up workbook shows that PLFC already has governance, public channels, Google tools, sponsor relationships and a clear growth constraint around volunteer and committee capacity."
)
journey_table = doc.add_table(rows=1, cols=4)
journey_table.alignment = WD_TABLE_ALIGNMENT.CENTER
journey_table.style = "Table Grid"
for cell, text in zip(journey_table.rows[0].cells, ["Signal", "What it means", "Backend response", "Public boundary"]):
    cell.text = text
for row in [
    ["Digital base exists", "Google Workspace is already the practical operating layer", "Use Sheets and Drive before heavier systems", "Do not publish account details"],
    ["Capacity is constrained", "People need clear ways to join, help and follow through", "Prioritise renewals, volunteers and action queues", "Do not expose member status"],
    ["Field future is strong", "Solunar, maps, media and catch data can grow over time", "Use public-safe exports and permissions", "No precise public GPS"],
    ["Broader ideas exist", "Apps, conservation and island ecosystems can connect later", "Keep schemas extensible", "Separate club-approved work from adjacent concepts"],
]:
    cells = journey_table.add_row().cells
    for cell, text in zip(cells, row):
        cell.text = text

doc.add_heading("6. Agent Handoff", level=1)
for item in [
    "Agent reads only authorised Drive, Sheet, Gmail or Calendar data.",
    "Agent removes precise GPS, private contacts, emergency details and unapproved media.",
    "Agent writes small public-safe JSON exports for website visualisations.",
    "Export records source, time, reviewer and public/private status.",
]:
    doc.add_paragraph(item, style="List Number")

doc.add_heading("7. Committee Approval Checklist", level=1)
for item in [
    "Approve Drive folder names and access groups.",
    "Approve Sheet columns and who can edit each sheet.",
    "Approve Calendar event types and invite rules.",
    "Approve media and location privacy wording.",
    "Approve which agent exports may update the public site.",
]:
    doc.add_paragraph(item, style="List Bullet")

footer = section.footer.paragraphs[0]
footer.text = "PLFC backend governance starter pack - prototype only"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(TEMPLATES / "PLFC Backend Governance Starter Pack.docx")
